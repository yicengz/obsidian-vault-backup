#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
从 ~/Desktop/无人知晓/*_原文.docx / *_导读.docx 生成中间 md
(说话人真名标注 + 主持人错字修正),写入 vault 无人知晓/ 目录。
之后由 merge_podcast.py 合并。
"""
import docx, re, os, glob, sys

SRC = os.path.expanduser("~/Desktop/无人知晓")
DST = "/Users/bytedance/Library/Mobile Documents/iCloud~md~obsidian/Documents/yiceng/raw/podcast/无人知晓"

MENG = "孟岩"
# 每期:host, guest(None=独白), 说话人标签->真名
def std(g):   return (MENG, g, {"发言人1": MENG, "发言人2": g})
def solo():   return (MENG, None, {"发言人1": MENG, "发言人": MENG})

CONFIG = {
    "E02": std("陈嘉禾"), "E03": std("张奔斗"), "E05": std("简七"),
    "E06": std("曹名长"), "E07": std("吴海燕"), "E08": std("方三文"),
    "E12": std("池建强"), "E14": std("爱哲"),  "E16": std("少楠"),
    "E17": std("张寒"),   "E18": std("陈鹏"),  "E22": std("老六"),
    "E24": std("孙方"),   "E25": std("季凯帆"), "E26": std("吴鲁加"),
    "E32": std("成庆"),   "E34": std("顾中一"), "E36": std("周奇墨"),
    "E38": std("陈行甲"), "E39": std("重轻"),  "E41": std("阿娇"),
    "E45": std("李继刚"),
    # 反转 / 串台
    "E04": ("孟岩", "张潇雨", {"发言人1": "张潇雨", "发言人2": MENG}),
    "E11": ("刘飞", "孟岩",   {"发言人1": MENG, "发言人2": "刘飞"}),
    "E13": ("Zara", "孟岩",   {"发言人1": MENG, "发言人2": "Zara"}),
    "E27": ("厚望", "孟岩",   {"发言人1": "厚望", "发言人2": MENG}),
    "E28": ("Steve", "孟岩",  {"发言人1": "Steve", "发言人2": MENG}),
    # 独白
    "E10": solo(), "E15": solo(), "E19": solo(), "E20": solo(), "E21": solo(),
    "E23": solo(), "E30": solo(), "E31": solo(), "E33": solo(), "E35": solo(),
    "E37": solo(), "E40": solo(),
}

def title_names(title):
    # "E43 A、B对话C:..." -> (["A","B"], "C")
    core = re.sub(r"^E\d+\s+", "", title).split("：")[0].split(":")[0]
    left, _, right = core.partition("对话")
    hosts = [x for x in re.split(r"[、,，]", left) if x]
    return hosts, right.strip()

def e43(title):  # 张潇雨、孟岩 对话 许哲;发言人1=孟岩,发言人2=张潇雨,发言人3=嘉宾
    hosts, guest = title_names(title)
    zhang = hosts[0]
    return (f"{zhang}、{MENG}", guest,
            {"发言人1": MENG, "发言人2": zhang, "发言人3": guest})

def e44(title):  # X 对话 孟岩;发言人1=孟岩(被访),发言人2=X(主持)
    hosts, _ = title_names(title)
    x = hosts[0]
    return (x, MENG, {"发言人1": MENG, "发言人2": x})

CONFIG["E43"] = e43
CONFIG["E44"] = e44

ALIAS = ["梦妍", "梦言", "梦魇", "孟妍"]  # 主持人孟岩的音转文错字(不含"莫言"以免误伤作家)

def fix(text):
    for a in ALIAS:
        text = text.replace(a, MENG)
    return text

def repl_labels(text, smap):
    # 先替长标签,再替裸"发言人"
    for k in sorted(smap, key=len, reverse=True):
        text = text.replace(k, smap[k])
    if "发言人" in text:  # 裸标签兜底
        text = text.replace("发言人", smap.get("发言人1", MENG))
    return fix(text)

def parse_date(s):
    m = re.search(r"(\d{4})年(\d{2})月(\d{2})日", s)
    return f"{m.group(1)}-{m.group(2)}-{m.group(3)}" if m else "2026-07-12"

TS_RE = re.compile(r"^(发言人\d*)\s+(\d{1,2}:\d{2}(?::\d{2})?)\s*$")

def header_block(host, guest):
    lines = [f"> 主播：{host}"]
    if guest:
        lines.append(f"> 嘉宾：{guest}")
    return lines

def gen_yuanwen(path, title, host, guest, smap):
    d = docx.Document(path)
    paras = d.paragraphs
    date = parse_date(paras[1].text if len(paras) > 1 else "")
    out = ["---", f"title: {title}", "podcast: 无人知晓",
           f"transcribed_at: {date}", "---", "", f"# {title}", ""]
    out += header_block(host, guest)
    out += ["", "---", "", "## 转录正文", ""]
    cur = None; ts = None
    for p in paras:
        t = p.text.strip()
        m = TS_RE.match(t)
        if m:
            cur = smap.get(m.group(1)) or smap.get("发言人1") or MENG
            ts = m.group(2)
            continue
        if cur and t:
            out.append(f"**【{cur} @ {ts}】** {fix(t)}")
            out.append("")
            cur = None
    return "\n".join(out) + "\n"

def gen_daodu(path, title, host, guest, smap):
    d = docx.Document(path)
    paras = d.paragraphs
    date = parse_date(paras[1].text if len(paras) > 1 else "")
    out = ["---", f"title: {title} · 导读", "podcast: 无人知晓",
           f"generated_at: {date}", "---", "", f"# {title} · 导读", ""]
    out += header_block(host, guest)
    out += ["", "---", ""]
    cur_section = None
    for p in paras[2:]:
        st = p.style.name
        t = p.text.strip()
        if st == "大标题":
            continue
        if st == "一级标题":
            out += ["", f"## {t}", ""]
            cur_section = t
        elif st == "二级正文":
            if cur_section == "章节速览":
                out.append(f"- **{repl_labels(t, smap)}**")
            else:
                out.append(repl_labels(t, smap))
            out.append("")
        else:  # Normal
            if not t:
                continue
            out.append(repl_labels(t, smap))
            out.append("")
    text = "\n".join(out)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text + "\n"

def main():
    eps = sys.argv[1:] or sorted(CONFIG.keys())
    done = []
    for ep in eps:
        if ep not in CONFIG:
            print(f"跳过 {ep}(无配置)"); continue
        yw = glob.glob(os.path.join(SRC, f"{ep} *_原文.docx"))
        dd = glob.glob(os.path.join(SRC, f"{ep} *_导读.docx"))
        if not yw or not dd:
            print(f"⚠️ {ep} 缺 docx: 原文={yw} 导读={dd}"); continue
        yw, dd = yw[0], dd[0]
        title = os.path.basename(yw).replace("_原文.docx", "")
        cfg = CONFIG[ep]
        host, guest, smap = cfg(title) if callable(cfg) else cfg
        with open(os.path.join(DST, f"{title}_原文.md"), "w", encoding="utf-8") as f:
            f.write(gen_yuanwen(yw, title, host, guest, smap))
        with open(os.path.join(DST, f"{title}_导读.md"), "w", encoding="utf-8") as f:
            f.write(gen_daodu(dd, title, host, guest, smap))
        print(f"✅ {ep} → 中间 md 已生成({host} / {guest})")
        done.append(ep)
    print(f"\n共生成 {len(done)} 期: {' '.join(done)}")

if __name__ == "__main__":
    main()
